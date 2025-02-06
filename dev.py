import streamlit as st
import base64
from googletrans import Translator
from langchain.chains import RetrievalQA
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.callbacks.manager import CallbackManager
from langchain.llms import Ollama
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain_community.vectorstores import Chroma
from langchain.embeddings.ollama import OllamaEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import pandas as pd
import os
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from database import get_mongo_client, get_database, get_collection1, get_collection2,get_collection3,get_collection_chat_process, save_chat_to_db
from hashlib import sha256
from datetime import datetime

logo_path = "C:/Users/DEV-037/Desktop/test-mesop/images/logo.png"


with open("C:/Users/DEV-037/Desktop/test-mesop/images/logo.png", "rb") as image_file:
    logo_base64 = base64.b64encode(image_file.read()).decode("utf-8")


st.markdown(f"""
    <style>
    .sidebar .sidebar-content {{
        background-color: #f0f2f6;
        border-radius: 10px;
        padding: 10px;
    }}
    .sidebar .sidebar-content h1 {{
        font-size: 24px;
        color: #1f77b4;
        font-weight: bold;
    }}
    .stTextInput>div>div>input {{
        border-radius: 10px;
        border: 1px solid #ddd;
        padding: 12px;
        font-size: 16px;
    }}
    .stButton>button {{
        background-color: #1f77b4;
        color: white;
        border: none;
        padding: 10px 20px;
        border-radius: 10px;
        font-size: 16px;
        font-weight: bold;
    }}
    .stButton>button:hover {{
        background-color: #155a8a;
        box-shadow: 0px 4px 6px rgba(0, 0, 0, 0.2);
    }}
    .chat-container {{
        border: 1px solid #ddd;
        border-radius: 10px;
        padding: 15px;
        background-color: #f9f9f9;
        margin-bottom: 10px;
        box-shadow: 0px 4px 6px rgba(0, 0, 0, 0.1);
    }}
    .message {{
        margin-bottom: 10px;
        font-size: 16px;
    }}
    .message.user {{
        color: #1f77b4;
        font-weight: bold;
    }}
    .message.assistant {{
        color: #e75f5f;
        font-weight: bold;
    }}
    .message.translated {{
        font-style: italic;
        color: #6c757d;
        font-size: 14px;
    }}
    .stDateInput>div>div {{
        font-size: 16px;
        padding: 10px;
        border-radius: 10px;
        border: 1px solid #ddd;
        background-color: #f0f2f6;
    }}
    .stContainer {{
        padding: 15px;
        background-color: #ffffff;
        border-radius: 10px;
        box-shadow: 0px 4px 6px rgba(0, 0, 0, 0.1);
    }}
    .small-button {{
        font-size: 12px;
        padding: 5px 10px;
        margin: 5px 0;
        border-radius: 5px;
    }}
    .center-text {{
        text-align: center;
        font-size: 50px;
        font-weight: bold;
        color: #333333;
        margin-top: 20px;
        margin-bottom: 20px;
        line-height: 1.2;
    }}
    .center-text p {{
        margin: 0;
    }}
    .bottom-logo {{
       
        display:flex;
        justify-content:center;
        align-items:center;
        text-align: center;
        font-size: 14px;
        color: #333333;
        position:fixed;
        bottom:0 !important;
        right:10% !important;
    }}
    .bottom-logo span{{
        font-weight:bold !important;
        
    }}
      .bottom-logo img{{
        margin-bottom:10px;
        
    }}
  
    </style>

    <div class="bottom-logo">
        <span class="powered-by">Powered by</span>
        <img src="data:image/png;base64,{logo_base64}" alt="Logo"/>
    </div>
""", unsafe_allow_html=True)


# Mongo setup
client = get_mongo_client()
db = get_database(client)
collection = get_collection1(db)
credentials_collection = get_collection2(db)
knowledge_base = get_collection3(db)
processed_files_collection = get_collection_chat_process(db)


def check_credentials(username, password):
    hashed_password = sha256(password.encode()).hexdigest()
    user = credentials_collection.find_one({"username": username, "password": hashed_password})
    return user is not None

def create_user(username, password):
    hashed_password = sha256(password.encode()).hexdigest()
    credentials_collection.insert_one({"username": username, "password": hashed_password})

def list_models(model_dir):
    return [model for model in os.listdir(model_dir) if os.path.isdir(os.path.join(model_dir, model))]

def is_admin(username, password):
    return username == "admin" and password == "123"

# Initialize session state if not already present
if 'template' not in st.session_state:
    st.session_state.template = """You are a knowledgeable chatbot, here to help with questions of the user. Your tone should be professional and informative.

    Context: {context}
    History: {history}

    User: {question}
    Chatbot:"""
if 'prompt' not in st.session_state:
    st.session_state.prompt = PromptTemplate(
        input_variables=["history", "context", "question"],
        template=st.session_state.template,
    )
if 'memory' not in st.session_state:
    st.session_state.memory = ConversationBufferMemory(
        memory_key="history",
        return_messages=True,
        input_key="question")

if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

if 'processed_files' not in st.session_state:
    st.session_state.processed_files = {}

MODEL_DIR = "C:/Users/DEV-037/.ollama/models/manifests/registry.ollama.ai/library"
FILE_DIR = 'files'
DATA_DIR = 'data'
translator = Translator()

os.makedirs(FILE_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)

available_models = list_models(MODEL_DIR)

# Session management UI
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'show_signup' not in st.session_state:
    st.session_state.show_signup = False
if 'is_admin' not in st.session_state:
    st.session_state.is_admin = False

# Get Knowledge basis from db    
def get_categories_from_db():
    categories = [doc['category_name'] for doc in db['knowledge_base'].find()]
    return categories

# Add Knowledge basis from db
def add_category_to_db(category_name):
    if category_name:
        db['knowledge_base'].insert_one({"category_name": category_name})
# Login
def login():
    st.sidebar.markdown('<div class="header">Login</div>', unsafe_allow_html=True)
    username = st.sidebar.text_input("Username", key="login_username", help="Enter your username")
    password = st.sidebar.text_input("Password", type="password", key="login_password", help="Enter your password")
    
    if st.sidebar.button("Login", key="login_btn", help="Click to login", use_container_width=True):
        if not username or not password:
            st.sidebar.error("Please enter both username and password.")
        elif is_admin(username, password):
            st.session_state.logged_in = True
            st.session_state.username = username
            st.session_state.is_admin = True
            reset_session_state()
            st.sidebar.success("Admin login successful!")
            st.rerun()
            
        elif check_credentials(username, password):
            st.session_state.logged_in = True
            st.session_state.username = username
            reset_session_state() 
            st.sidebar.success("Login successful!")
            st.rerun()
            
        else:
            st.sidebar.error("Invalid username or password")
            
# Session State Reset           
def reset_session_state():
    st.session_state.chat_history = [] 
    st.session_state.processed_files = {} 
    st.session_state.vectorstore = None  
    st.session_state.qa_chain = None  
            
# Sign-up
def signup():
    st.sidebar.markdown('<div class="header">Sign Up</div>', unsafe_allow_html=True)
    username = st.sidebar.text_input("Username", key="signup_username", help="Choose a new username")
    password = st.sidebar.text_input("Password", type="password", key="signup_password", help="Set your password")
    confirm_password = st.sidebar.text_input("Confirm Password", type="password", key="confirm_password", help="Re-enter your password")
    
    if st.sidebar.button("Sign Up", key="signup_btn", help="Create a new account", use_container_width=True):
        if not username or not password or not confirm_password:
            st.sidebar.error("Please enter both username and password.")
        elif password  != confirm_password:
            st.sidebar.error("Passwords do not match. Please try again.")   
        elif check_credentials(username, password):
            st.sidebar.error("Username already exists. Please choose a different username.")
        else:
            create_user(username, password)
            st.sidebar.success("User created successfully!")
            st.session_state.logged_in = True  
            st.session_state.username = username
            st.rerun() 
            
def extract_text_from_pdf(file_path):
        text = ""
        with open(file_path, "rb") as file:
            reader = PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text() + "\n"
        return text

def extract_text_from_docx(file_path):
    text = ""
    doc = Document(file_path)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def adjust_chunk_parameters(text):
    total_length = len(text)
    avg_sentence_length = sum(len(sentence) for sentence in text.split(".")) / max(len(text.split(".")), 1)

    if total_length < 5000:
        chunk_size = 600
        chunk_overlap = 100
    elif total_length < 20000:
        chunk_size = 800
        chunk_overlap = 200
    else:
        chunk_size = 1000
        chunk_overlap = 300

    if avg_sentence_length > 100:
        chunk_size += 500
        chunk_overlap += 100

    return chunk_size, chunk_overlap 

def check_file_in_db(file_name, file_size, model_name):
    return processed_files_collection.find_one({
        "file_name": file_name,
        "file_size": file_size,
        "model_name": model_name
    })

def add_file_to_db(file_name, file_size, model_name, file_path, vectorstore_path):
    processed_files_collection.insert_one({
        "file_name": file_name,
        "file_size": file_size,
        "model_name": model_name,
        "file_path": file_path,
        "vectorstore_path": vectorstore_path
    })


# Logic to toggle between login and signup form
if not st.session_state.logged_in:
    st.markdown('<div class="center-text"><p>Please sign in or sign up to proceed further</p></div>', unsafe_allow_html=True)
    if st.sidebar.button("Sign In", key="signin", help="Already have an account?", use_container_width=True):
        st.session_state.show_signup = False
    
    if st.sidebar.button("Sign Up", key="signup", help="Create a new account", use_container_width=True):
        st.session_state.show_signup = True

    if st.session_state.show_signup:
        signup()
    else:
        login()
else:
    st.sidebar.header(f"Welcome {st.session_state.username.capitalize()}")
    
    
    # Admin functionality ..get chats by username,date and category
    if st.session_state.is_admin:
        users = list(credentials_collection.find({}))
        usernames = [user["username"] for user in users]
        
        # Dropdown to select user
        usernames.insert(0, "Click here to view users")
        selected_user = st.sidebar.selectbox("Select User", options=usernames, index=0)
        
        # Dropdown to select category
        categories = get_categories_from_db()
        categories.insert(0, "Click here to view categories")
        selected_category = st.sidebar.selectbox("Select Category", options=categories, index=0)

        if selected_user and selected_category:
            date_input = st.sidebar.date_input("📅 Enter Date to view chats", min_value=None)
           
            def retrieve_data_by_user_and_date_and_category(user, date,category):
                formatted_date = date.strftime('%d-%m-%Y')
                query = {"username": user, "timestamp": formatted_date,"category": category}
                return list(collection.find(query))
    
            if st.sidebar.button("History", key="submit_button", use_container_width=True):
                if date_input:
                    formatted_date = date_input.strftime('%d-%m-%Y')
                    data = retrieve_data_by_user_and_date_and_category(selected_user, date_input,selected_category)

                    if data:
                        st.sidebar.subheader(f"Chats for {selected_user} in {selected_category} on {formatted_date}")
                        for record in data:
                            st.sidebar.write(f"**💬 User Query:** {record['user_query']}")
                            st.sidebar.write(f"**🤖 Assistant Response:** {record['assistant_response']}")
                            # st.sidebar.write(f"**🌐 Translated Response:** {record['translated_response']}")
                            st.sidebar.write(f"**🕒 Timestamp:** {record['timestamp']}")
                            st.sidebar.write(f"**⚙️ Model:** {record['model_name']}")
                            st.sidebar.write("---")
                    else:
                        st.sidebar.write(f"No records found for {selected_user} in {selected_category} on {formatted_date}.")
                else:
                    st.sidebar.write("Please select a date.")
                    
                         
    # Logout button for both users and admins
    if st.sidebar.button("Logout", key="logout_btn", help="Click to logout", use_container_width=True):
        # Reset session state
        st.session_state.logged_in = False
        st.session_state.username = ""
        st.session_state.is_admin = False
        st.session_state.show_signup = False 
        st.sidebar.success("You have been logged out successfully!")
        st.rerun() 
        
    # chat               
    if 'selected_model' not in st.session_state:
        st.session_state.selected_model = None
    if 'llm' not in st.session_state:
        st.session_state.llm = None
    if 'vectorstore' not in st.session_state:
        st.session_state.vectorstore = None
    if 'qa_chain' not in st.session_state:
        st.session_state.qa_chain = None

    # Model selection
    st.sidebar.header("Model Selection")
    # available_models = ["Mistral", "Llama 3.1", "Other Models..."]
    st.session_state.selected_model = st.sidebar.selectbox("Select Model", available_models, index=0)

    if st.session_state.selected_model:
        # Configure the model
        st.session_state.llm = Ollama(
            base_url="http://localhost:11434",
            model=st.session_state.selected_model,
            verbose=True,
            callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]),
        )
        st.session_state.embedding_model = OllamaEmbeddings(model=st.session_state.selected_model)

        # File upload and processing
        uploaded_files = st.sidebar.file_uploader(
            "Upload your files",
            type=['xlsx', 'xls', 'pdf', 'docx'],
            accept_multiple_files=True
        )

        if uploaded_files:
            for uploaded_file in uploaded_files:
                file_key = f"{uploaded_file.name}_{uploaded_file.size}"
                file_path = os.path.join("data", uploaded_file.name)
                vectorstore_path = os.path.join("data", f"{uploaded_file.name}_vectorstore")

                # Check MongoDB if file is processed with this model
                file_in_db = check_file_in_db(uploaded_file.name, uploaded_file.size, st.session_state.selected_model)
                if file_in_db:
                    st.sidebar.warning(f"⚠️ File '{uploaded_file.name}' already processed. You can ask questions.")
                    st.session_state.vectorstore = Chroma(persist_directory=file_in_db["vectorstore_path"], embedding_function=st.session_state.embedding_model)
                else:
                    # Save file and process it
                    with open(file_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())

                    # Process file content based on type
                    if uploaded_file.name.endswith('.xlsx') or uploaded_file.name.endswith('.xls'):
                        xls = pd.ExcelFile(file_path)
                        all_content = "\n".join([df.to_dict(orient='records') for df in [pd.read_excel(xls, sheet) for sheet in xls.sheet_names]])
                    elif uploaded_file.name.endswith('.pdf'):
                        all_content = extract_text_from_pdf(file_path)
                    elif uploaded_file.name.endswith('.docx'):
                        all_content = extract_text_from_docx(file_path)
                    else:
                        st.sidebar.warning(f"Unsupported file type: {uploaded_file.name}")
                        continue

                    # Chunk and embed the document
                    chunk_size, chunk_overlap = adjust_chunk_parameters(all_content)
                    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap, length_function=len)
                    splits = text_splitter.split_text(all_content)
                    documents = [Document(page_content=split) for split in splits]
                    ids = [f"{uploaded_file.name}_chunk_{i}" for i in range(len(documents))]

                    # Create and persist vector store
                    st.session_state.vectorstore = Chroma.from_documents(documents=documents, embedding=st.session_state.embedding_model,
                                                                        persist_directory=vectorstore_path, ids=ids)
                    st.session_state.vectorstore.persist()
                    
                    # Store processed file details in MongoDB
                    add_file_to_db(uploaded_file.name, uploaded_file.size, st.session_state.selected_model, file_path, vectorstore_path)
                    st.sidebar.success(f"File '{uploaded_file.name}' processed successfully with model '{st.session_state.selected_model}'!")

            # Initialize retriever and QA chain if model matches
            if st.session_state.vectorstore:
                st.session_state.retriever = st.session_state.vectorstore.as_retriever()
                st.session_state.qa_chain = RetrievalQA.from_chain_type(
                    llm=st.session_state.llm,
                    chain_type='stuff',
                    retriever=st.session_state.retriever,
                    verbose=True,
                    chain_type_kwargs={
                        "verbose": True,
                        "prompt": st.session_state.prompt,
                        "memory": st.session_state.memory,
                    }
                )
    else:
        st.sidebar.error("Please select a model before uploading files.")
    
    # Layout containers
    chat_container = st.container()
    
    # Display Chat History
    with chat_container:
        st.subheader("💬 Chat")
        for message in st.session_state.chat_history:
            if message['role'] == 'user':
                st.markdown(f"<div class='chat-container'><p class='message user'> 👤 {message['message']}</p></div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div class='chat-container'><p class='message assistant'> 🧠 {message['message']}</p></div>", unsafe_allow_html=True)
                
    # Input Field at the Bottom
    input_container = st.container()
    with input_container:
        st.write("---")
        user_input = st.text_input(
            "💬 Ask a question:",
            key="user_input",
            placeholder="Type your question here..."
        )
        
    #Translation
    def safe_translate(text, dest_lang):
        try:
            translated = translator.translate(text, dest=dest_lang)
            return translated.text
        except Exception as e:
            st.error(f"Translation failed: {str(e)}")
            return text
        
    # Overall mongodb Saving Function
    def save_to_mongo(user_input, response_text, translated_response):
        current_time = datetime.now().strftime('%d-%m-%Y')
        chat_data = {
            "username": st.session_state.username,  
            "user_query": user_input,
            "assistant_response": response_text,
            "translated_response": translated_response,
            "timestamp": current_time,
            "model_name": st.session_state.selected_model,
            "category": st.session_state.get('selected_category', None) 
        }
        save_chat_to_db(collection, chat_data)

    if st.button("Submit"):
        if user_input:
            user_message = {"role": "user", "message": user_input}
            st.session_state.chat_history.append(user_message)
            st.markdown(f"<div class='chat-container'><p class='message user'>{user_input}</p></div>", unsafe_allow_html=True)

            with st.spinner("Assistant is typing..."):
                if st.session_state.qa_chain:
                    response = st.session_state.qa_chain({"query": user_input})
                    response_text = response['result']

                    translated_response = safe_translate(response_text, 'ta')
                    
                    # MongoDB
                    save_to_mongo(user_input, response_text, translated_response)

                    assistant_message = {"role": "assistant", "message": response_text, "translated_message": translated_response}
                    
                    st.session_state.chat_history.append(assistant_message)

                    st.markdown(f"<div class='chat-container'><p class='message assistant'>{response_text}</p></div>", unsafe_allow_html=True)
                else:
                    st.error("QA chain not initialized. Please upload a document to start the chat.")
        else:
            st.warning("Please enter a question before submitting.")
            
            
            
            
            
            
            
            
            
    #chat retrieval for both admin and users....and viewable in both screens
    # if st.session_state.logged_in:
    #     date_input = st.sidebar.date_input("📅 Enter Date to view previous chats", min_value=None,key="overall")

    #     def retrieve_data_by_date(date):
    #         formatted_date = date.strftime('%d-%m-%Y')
    #         query = {"timestamp": formatted_date}
    #         return list(collection.find(query))

    #     if st.sidebar.button("History", key="history_retrival"):
    #         if date_input:
    #             formatted_date = date_input.strftime('%d-%m-%Y')
    #             data = retrieve_data_by_date(date_input)
                
    #             if data:
    #                 st.sidebar.subheader(f"Data for {formatted_date}")
    #                 for record in data:
    #                     st.sidebar.write(f"**💬 User Query:** {record['user_query']}")
    #                     st.sidebar.write(f"**🤖 Assistant Response:** {record['assistant_response']}")
    #                     st.sidebar.write(f"**🌐 Translated Response:** {record['translated_response']}")
    #                     st.sidebar.write(f"**🕒 Timestamp:** {record['timestamp']}")
    #                     st.sidebar.write(f"**⚙️ Model:** {record['model_name']}")
    #                     st.sidebar.write("---")
    #             else:
    #                 st.sidebar.write(f"No records found for {formatted_date}.")
    #         else:
    #             st.sidebar.write("Please select a date.")
    # else:
    #     st.sidebar.warning("You need to be logged in to view chat history.")
