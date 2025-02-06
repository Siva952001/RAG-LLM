import os
import pandas as pd
# import numpy as np
import streamlit as st
from PyPDF2 import PdfReader
from statistics import mean
from docx import Document
from langchain.schema import Document as LangchainDocument
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter

FILE_DIR = 'files'

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, "rb") as file:
        reader = PdfReader(file)
        for page_num in range(len(reader.pages)):
            text += reader.pages[page_num].extract_text() + "\n"
    return text

def extract_text_from_docx(file_path):
    text = ""
    doc = Document(file_path)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Chunking function
def calculate_chunk_parameters(text):
    total_length = len(text)
    sentences = text.split(".")
    num_sentences = len(sentences)
    avg_sentence_length = mean(len(sentence.strip()) for sentence in sentences if sentence.strip())
    words = text.split()
    num_words = len(words)
    avg_word_length = mean(len(word) for word in words) if num_words > 0 else 0
    text_density = num_words / num_sentences if num_sentences > 0 else 0

    # Base chunk size and overlap
    chunk_size = 600
    chunk_overlap = 100

    # Adjust chunk size based on text length
    if total_length < 5000:
        chunk_size += 0
    elif total_length < 20000:
        chunk_size += 200
    else:
        chunk_size += 400

    # Adjust chunk size based on average sentence length
    if avg_sentence_length > 100:
        chunk_size += 300
    elif avg_sentence_length < 20:
        chunk_size -= 100

    # Adjust chunk size based on average word length
    if avg_word_length > 6:
        chunk_size += 100
    elif avg_word_length < 4:
        chunk_size -= 50

    # Adjust chunk size based on text density
    if text_density > 20:
        chunk_size += 200
    elif text_density < 10:
        chunk_size -= 100

    # Overlap adjustments (dynamic scaling)
    chunk_overlap = max(100, int(chunk_size * 0.2))

    # Ensure chunk_size and chunk_overlap are reasonable
    chunk_size = max(chunk_size, 300)  
    chunk_overlap = min(chunk_overlap, chunk_size // 2)  

    return chunk_size, chunk_overlap

def process_uploaded_files(files, directory):
    if files:
        for file in files:
            file_key = f"{file.name}_{file.size}"
            file_path = os.path.join(FILE_DIR, file.name)
            vectorstore_path = os.path.join(directory, f"{file.name}_vectorstore")

            try:
                # Check if file is already processed with the current embedding model
                if file_key in st.session_state.processed_files:
                    st.sidebar.warning(f"⚠️ File '{file.name}' already exists. You can ask questions.")
                elif os.path.exists(vectorstore_path):
                    try:
                        # Check embedding dimensionality before loading
                        st.session_state.vectorstore = Chroma(
                            persist_directory=vectorstore_path,
                            embedding_function=st.session_state.embedding_model
                        )
                        st.session_state.retriever = st.session_state.vectorstore.as_retriever()
                        st.session_state.processed_files[file_key] = True
                    except ValueError:
                        st.sidebar.warning("⚠️ File could not be loaded due to internal issues. Please try again.")
                        continue
                else:
                    with open(file_path, "wb") as f:
                        f.write(file.getvalue())

                    with st.spinner("Training started..."):
                        try:
                            # Extract content based on file type
                            if file.name.endswith('.xlsx') or file.name.endswith('.xls'):
                                xls = pd.ExcelFile(file_path)
                                all_content = "\n".join(
                                    ["\n".join(map(str, pd.read_excel(xls, sheet).to_dict(orient='records')))
                                     for sheet in xls.sheet_names]
                                )
                            elif file.name.endswith('.pdf'):
                                all_content = extract_text_from_pdf(file_path)
                            elif file.name.endswith('.docx'):
                                all_content = extract_text_from_docx(file_path)
                            else:
                                st.sidebar.warning(f"Unsupported file type: {file.name}")
                                continue

                            # Determine chunking parameters
                            chunk_size, chunk_overlap = calculate_chunk_parameters(all_content)
                            text_splitter = RecursiveCharacterTextSplitter(
                                chunk_size=chunk_size,
                                chunk_overlap=chunk_overlap,
                                length_function=len
                            )

                            st.sidebar.write(f"### Please wait while '{file.name}' is being processed")
                            st.sidebar.write(f"- **Chunk Size**: {chunk_size}")
                            st.sidebar.write(f"- **Chunk Overlap**: {chunk_overlap}")

                            # Split text and create documents
                            splits = text_splitter.split_text(all_content)
                            documents = [LangchainDocument(page_content=split) for split in splits]

                            # Generate unique IDs for each chunk
                            ids = [f"{file.name}_chunk_{i}" for i in range(len(documents))]

                            # Create or update vector store
                            st.session_state.vectorstore = Chroma.from_documents(
                                documents=documents,
                                embedding=st.session_state.embedding_model,
                                persist_directory=vectorstore_path,
                                ids=ids
                            )
                            st.session_state.vectorstore.persist()
                            st.session_state.processed_files[file_key] = True
                            st.session_state.retriever = st.session_state.vectorstore.as_retriever()
                            st.sidebar.success(f"File '{file.name}' processed successfully!")
                        except Exception:
                            st.sidebar.error("⚠️ Unable to process the file due to internal issues. Please try again later.")
                            continue
            except Exception:
                st.sidebar.error("⚠️ An unexpected error occurred. Please contact support or try again.")