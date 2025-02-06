# Imports
import streamlit as st
from streamlit_option_menu import option_menu
import base64
from langchain_community.vectorstores import Chroma
# from googletrans import Translator
from langchain.chains import RetrievalQA
from langchain.retrievers import EnsembleRetriever
# from statistics import mean
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
# from sklearn.decomposition import PCA
# from sklearn.preprocessing import StandardScaler
from langchain.callbacks.manager import CallbackManager
# import numpy as np
from langchain.llms import Ollama
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain.embeddings.ollama import OllamaEmbeddings
import pandas as pd
import re
import os
from database import get_mongo_client, get_database, get_collection1, get_collection2,get_collection3, save_chat_to_db
from hashlib import sha256
from datetime import datetime
from chat_retrieval import load_chat_retrieval,get_categories_from_db
from file_process import process_uploaded_files
import os
import pandas as pd
# import numpy as np
import streamlit as st
from PyPDF2 import PdfReader
from statistics import mean
from docx import Document
from langchain.schema import Document as LangchainDocument
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import Chroma
# import os
# import pandas as pd
# import numpy as np
# import streamlit as st
# from PyPDF2 import PdfReader
# from docx import Document
# from langchain.schema import Document as LangchainDocument
# from langchain_community.vectorstores import Chroma
# from langchain.text_splitter import RecursiveCharacterTextSplitter


logo_path = "images/logo.png"


with open("images/logo.png", "rb") as image_file:
    logo_base64 = base64.b64encode(image_file.read()).decode("utf-8")

st.markdown(f"""
    <style>
    .center-text {{
        text-align: center;
        font-size: 50px;
        font-weight: bold;
        color: #333333;
        margin-top: 20px;
        margin-bottom: 20px;
        line-height: 1.2;
    }}
    .center-text p {{
        margin: 0;
    }}
    .bottom-logo {{
       
        display:flex;
        justify-content:center;
        align-items:center;
        text-align: center;
        font-size: 14px;
        color: #333333;
        position:fixed;
        bottom:0 !important;
        right:10% !important;
    }}
    .bottom-logo span{{
        font-weight:bold !important;
        
    }}
      .bottom-logo img{{
        margin-bottom:10px;
        
    }}
  
    </style>

    <div class="bottom-logo">
        <span class="powered-by">Powered by</span>
        <img src="data:image/png;base64,{logo_base64}" alt="Logo"/>
    </div>
""", unsafe_allow_html=True)

# Mongo setup
client = get_mongo_client()
db = get_database(client)
collection = get_collection1(db)
credentials_collection = get_collection2(db)
knowledge_base = get_collection3(db)

def username_exists(username):
    user = credentials_collection.find_one({"username": username})
    return user is not None

# Check Credentials
def check_credentials(username, password):
    hashed_password = sha256(password.encode()).hexdigest()
    user = credentials_collection.find_one({"username": username, "password": hashed_password})
    return user is not None

# Create User 
def create_user(username, password):
    hashed_password = sha256(password.encode()).hexdigest()
    credentials_collection.insert_one({"username": username, "password": hashed_password})
 
# Models List   
def list_models(model_dir):
    return [model for model in os.listdir(model_dir) if os.path.isdir(os.path.join(model_dir, model))]

# Admin Credential
def is_admin(username, password):
    return username == "admin" and password == "123"

# Password validation
def check_password_strength(password):
    """Check if the password meets strength requirements."""
    if len(password) < 7:
        return False
    if not re.search(r"[A-Z]", password):  
        return False
    if not re.search(r"[a-z]", password):  
        return False
    if not re.search(r"[!@#$%^&*(),.?\":{}|<>]", password):  
        return False
    return True

# Overall mongodb Saving Function
def save_to_mongo(user_input, response_text):
        current_time = datetime.now().strftime('%d-%m-%Y')
        chat_data = {
            "username": st.session_state.username,  
            "user_query": user_input,
            "assistant_response": response_text,
            # "translated_response": translated_response,
            "timestamp": current_time,
            "model_name": st.session_state.selected_model,
            "category": st.session_state.get('selected_category', None) 
        }
        save_chat_to_db(collection, chat_data)
        
        
# Get User Chat History        
def get_user_chat_history(username, start_date=None, end_date=None):
    try:
        query = {"username": username}
        
        if start_date and end_date:
            formatted_start_date = start_date.strftime("%d-%m-%Y")
            formatted_end_date = end_date.strftime("%d-%m-%Y")
            query["timestamp"] = {"$gte": formatted_start_date, "$lte": formatted_end_date}
        
        chat_history = list(collection.find(query).sort("timestamp"))
        return chat_history
    except Exception as e:
        print(f"An error occurred while retrieving chat history: {e}")
        return []
    

# Initialize session state if not already present
if 'template' not in st.session_state:
    st.session_state.template = """You are a knowledgeable chatbot, here to help with questions of the user. Your tone should be professional and informative.

    Context: {context}
    History: {history}

    User: {question}
    Chatbot:"""
if 'prompt' not in st.session_state:
    st.session_state.prompt = PromptTemplate(
        input_variables=["history", "context", "question"],
        template=st.session_state.template,
    )
if 'memory' not in st.session_state:
    st.session_state.memory = ConversationBufferMemory(
        memory_key="history",
        return_messages=True,
        input_key="question")

if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

if 'processed_files' not in st.session_state:
    st.session_state.processed_files = {}

MODEL_DIR = "C:/Users/DEV-037/.ollama/models/manifests/registry.ollama.ai/library"
FILE_DIR = 'files'
DATA_DIR = 'data'
# translator = Translator()

os.makedirs(FILE_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)

available_models = list_models(MODEL_DIR)

# Session management UI
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'show_signup' not in st.session_state:
    st.session_state.show_signup = False
if 'is_admin' not in st.session_state:
    st.session_state.is_admin = False

     
# Login
def login():
    with st.form(key='login_form', clear_on_submit=True):
        username = st.text_input("Username", key="login_username", help="Enter your username")
        password = st.text_input("Password", type="password", key="login_password", help="Enter your password")
        submit_button = st.form_submit_button("Login", help="Click to login")
        
        if submit_button:
            if not username or not password:
                st.error("Please enter both username and password.")
            elif is_admin(username, password):
                st.session_state.logged_in = True
                st.session_state.username = username
                st.session_state.is_admin = True
                reset_session_state()
                st.success("Admin login successful!")
                st.rerun()
            elif check_credentials(username, password):
                st.session_state.logged_in = True
                st.session_state.username = username
                reset_session_state()
                st.success("Login successful!")
                st.rerun()
            else:
                st.error("Invalid username or password")
                
# Session State Reset           
def reset_session_state():
    st.session_state.chat_history = [] 
    st.session_state.processed_files = {} 
    st.session_state.vectorstore = None  
    st.session_state.qa_chain = None  
            
            
# Sign-up 
def signup():
    with st.form(key='signup_form', clear_on_submit=True):
        username = st.text_input("Username", key="signup_username", help="Choose a new username")
        password = st.text_input("Password", type="password", key="signup_password", help="Set your password")
        confirm_password = st.text_input("Confirm Password", type="password", key="confirm_password", help="Re-enter your password")   
        
        submit_button = st.form_submit_button("Sign Up", help="Create a new account")
        
        if submit_button:
            if not username or not password or not confirm_password:
                st.error("Please enter both username and password.")
            elif password != confirm_password:
                st.error("Passwords do not match. Please try again.")
            elif not check_password_strength(password):
                st.error("Password must be at least 7 characters long, contain at least one uppercase letter, one lowercase letter, and one special character.")
            elif username_exists(username):
                st.error("Username already exists. Please choose a different username.")
            else:
                create_user(username, password)
                st.success("User created successfully!")
                st.session_state.logged_in = True
                st.session_state.username = username
                st.rerun()

# Main logic to toggle between login and signup form
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in: 
    selected = option_menu(
        menu_title=None, 
        options=["Login", "Sign Up"],  
        icons=["box-arrow-in-right", "person-plus"], 
        menu_icon="cast",  
        default_index=0,  
        orientation="horizontal",  
        styles = {
    "container": {
        "padding": "0!important", 
        "background-color": "#ffffff",  
        "transition": "background-color 0.3s ease"  
    },
    "icon": {
        "color": "#007bff",  
        "font-size": "25px",
        "transition": "color 0.3s ease"  
    },
    "nav-link": {
        "font-size": "16px", 
        "text-align": "center", 
        "margin": "0px", 
        "--hover-color": "#f0f0f0",  
        "color": "#333",  
        "transition": "color 0.3s ease, background-color 0.3s ease"  
    },
    "nav-link-hover": {
        "background-color": "#e5e5e5",  
        "color": "#0056b3",  
        "transition": "background-color 0.3s ease, color 0.3s ease"  
    },
    "nav-link-selected": {
        "background-color": "#0056b3",  
        "color": "white", 
        "font-weight": "bold",  
        "transition": "background-color 0.3s ease, color 0.3s ease"  
    }
 }
)

    if selected == "Login":
        login()
    elif selected == "Sign Up":
        signup()
else:
    # st.sidebar.markdown("You are now logged in.")
    st.sidebar.markdown(
    f"""
    <div style="text-align: center; position:relative; bottom:88px ; right:12px; margin-bottom:16px; padding:15px;">
        <img src="data:image/png;base64,{logo_base64}" alt="Logo" style="width: 150px;", "line-height:5.6," />
    </div>
    """,
    unsafe_allow_html=True
)
    
    # Add category for admin  
    def add_category_to_db(category_name):
        if category_name:
            db['knowledge_base'].insert_one({"category_name": category_name})
            
    # Add New Category        
    def add_new_category():
        if st.session_state.is_admin:  
            new_category = st.sidebar.text_input("Enter New Category Name")
            if st.sidebar.button("Add Category") and new_category:
                if db['knowledge_base'].find_one({"category_name": new_category}):
                    st.sidebar.error("Category already exists.")
                else:
                    add_category_to_db(new_category)
                    st.sidebar.success(f"'{new_category}' added to categories.")
                    st.rerun()  
        else:
            st.sidebar.warning("You need admin privileges to add a new category.")
            
    # File Processing
    FILE_DIR = 'files'

    def extract_text_from_pdf(file_path):
        text = ""
        with open(file_path, "rb") as file:
            reader = PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text() + "\n"
        return text

    def extract_text_from_docx(file_path):
        text = ""
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    
    def extract_text_from_csv(file_path):
        try:
            df = pd.read_csv(file_path)
            return df.to_csv(index=False, sep="\t")  # Convert to a tab-separated string
        except Exception as e:
            return f"Error extracting text from CSV: {e}"
        
    def extract_text_from_txt(file_path):
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except Exception as e:
            return f"Error extracting text from TXT: {e}"

    # Chunking function
    def calculate_chunk_parameters(text):
        total_length = len(text)
        sentences = text.split(".")
        num_sentences = len(sentences)
        avg_sentence_length = mean(len(sentence.strip()) for sentence in sentences if sentence.strip())
        words = text.split()
        num_words = len(words)
        avg_word_length = mean(len(word) for word in words) if num_words > 0 else 0
        text_density = num_words / num_sentences if num_sentences > 0 else 0

        # Base chunk size and overlap
        chunk_size = 600
        chunk_overlap = 100

        # Adjust chunk size based on text length
        if total_length < 5000:
            chunk_size += 0
        elif total_length < 20000:
            chunk_size += 200
        else:
            chunk_size += 400

        # Adjust chunk size based on average sentence length
        if avg_sentence_length > 100:
            chunk_size += 300
        elif avg_sentence_length < 20:
            chunk_size -= 100

        # Adjust chunk size based on average word length
        if avg_word_length > 6:
            chunk_size += 100
        elif avg_word_length < 4:
            chunk_size -= 50

        # Adjust chunk size based on text density
        if text_density > 20:
            chunk_size += 200
        elif text_density < 10:
            chunk_size -= 100

        # Overlap adjustments (dynamic scaling)
        chunk_overlap = max(100, int(chunk_size * 0.2))

        # Ensure chunk_size and chunk_overlap are reasonable
        chunk_size = max(chunk_size, 300)  
        chunk_overlap = min(chunk_overlap, chunk_size // 2)  

        return chunk_size, chunk_overlap

    def process_uploaded_files(files, category, model_name):
        if files:
            # Create model-specific directory under the category
            model_dir = os.path.join(DATA_DIR, category, model_name)
            os.makedirs(model_dir, exist_ok=True)
            
            for file in files:
                vectorstore_path = os.path.join(model_dir, f"{file.name}_vectorstore")
                # Check only within the current model's directory
                if os.path.exists(vectorstore_path):
                    st.sidebar.warning(f"⚠️ The file '{file.name}' is already uploaded and processed for the model '{model_name}'.")
                    continue

                try:
                    # Save the uploaded file to the `FILE_DIR` for processing
                    file_path = os.path.join(FILE_DIR, file.name)
                    with open(file_path, "wb") as f:
                        f.write(file.getvalue())
                    
                    with st.spinner(f"Processing {file.name} using {model_name}..."):
                        # Extract content based on file type
                        if file.name.endswith('.xlsx') or file.name.endswith('.xls'):
                            xls = pd.ExcelFile(file_path)
                            all_content = "\n".join(
                                ["\n".join(map(str, pd.read_excel(xls, sheet).to_dict(orient='records')))
                                for sheet in xls.sheet_names]
                            )
                        elif file.name.endswith('.pdf'):
                            all_content = extract_text_from_pdf(file_path)
                        elif file.name.endswith('.docx'):
                            all_content = extract_text_from_docx(file_path)
                        elif file.name.endswith('.csv'):
                            all_content = extract_text_from_csv(file_path)
                        elif file.name.endswith('.txt'):
                            all_content = extract_text_from_txt(file_path)
                        else:
                            st.sidebar.warning(f"Unsupported file type: {file.name}")
                            continue

                        # Split text into chunks
                        chunk_size, chunk_overlap = calculate_chunk_parameters(all_content)
                        text_splitter = RecursiveCharacterTextSplitter(
                            chunk_size=chunk_size,
                            chunk_overlap=chunk_overlap,
                            length_function=len
                        )
                        splits = text_splitter.split_text(all_content)
                        documents = [LangchainDocument(page_content=split) for split in splits]
                        ids = [f"{file.name}_chunk_{i}" for i in range(len(documents))]

                        # Create and persist vectorstore in the model-specific directory
                        st.session_state.vectorstore = Chroma.from_documents(
                            documents=documents,
                            embedding=st.session_state.embedding_model,
                            persist_directory=vectorstore_path,
                            ids=ids
                        )
                        st.session_state.vectorstore.persist()
                        st.sidebar.success(f"✅ File {file.name} processed successfully for model {model_name}")
                except Exception as e:
                    st.sidebar.error(f"❌ Error processing file '{file.name}': {e}")
                    
    # Initiate Chat                
    def initiate_chat():
        try:
            # Step 1: Initialize session state for critical variables
            st.session_state.setdefault("selected_model", available_models[0])
            st.session_state.setdefault("llm", None)
            st.session_state.setdefault("embedding_model", None)
            st.session_state.setdefault("vectorstore", None)
            st.session_state.setdefault("qa_chain", None)
            st.session_state.setdefault("retrievers", [])
            st.session_state.setdefault("chat_history", [])
            st.session_state.setdefault("files_uploaded", False)
            st.session_state.setdefault("uploaded_file_names", [])
            st.session_state.setdefault("processed_files", {})
            
            # Step 2: Initialize or switch models
            selected_model = st.sidebar.selectbox(
                "Select Model",
                available_models,
                index=available_models.index(st.session_state.selected_model)
            )

            if st.session_state.selected_model != selected_model or st.session_state.llm is None:
                try:
                    st.session_state.selected_model = selected_model
                    st.session_state.llm = Ollama(
                        base_url="http://localhost:11434",
                        model=selected_model,
                        verbose=True,
                        callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]),
                    )
                    st.session_state.embedding_model = OllamaEmbeddings(model=selected_model)
                    st.session_state.retrievers = []
                    st.toast(f"Switched to model: {selected_model}")
                except Exception as e:
                    st.error("Failed to initialize the selected model. Please check the server and model configurations.")
                    return

            # Step 3: Select or create a category
            try:
                category_options = get_categories_from_db()
            except Exception:
                st.warning("Could not fetch categories from the database. Using default options.")
                category_options = []

            if st.session_state.is_admin:
                category_options.append("+ Create New Category")

            selected_category = st.sidebar.selectbox("Select Domain/Project", category_options)
            st.session_state.selected_category = selected_category

            if selected_category == "+ Create New Category":
                try:
                    add_new_category()
                except Exception:
                    st.error("An error occurred while creating a new category. Please try again.")
                    return

            # Step 4: Load vectorstores if available
            if selected_category and selected_category != "+ Create New Category":
                model_dir = os.path.join(DATA_DIR, selected_category, st.session_state.selected_model)
                try:
                    if os.path.exists(model_dir):
                        vectorstore_files = [
                            f for f in os.listdir(model_dir) if f.endswith("_vectorstore")
                        ]
                    else:
                        vectorstore_files = []
                except Exception:
                    st.warning("Could not list files in the selected category. Please check your permissions.")
                    vectorstore_files = []

                if vectorstore_files:
                    selected_existing_files = st.sidebar.multiselect(
                        "Select Knowledge Base:",
                        vectorstore_files,
                    )
                    st.session_state.retrievers = []
                    for file in selected_existing_files:
                        try:
                            vectorstore_path = os.path.join(model_dir, file)
                            vectorstore = Chroma(
                                persist_directory=vectorstore_path,
                                embedding_function=st.session_state.embedding_model
                            )
                            st.session_state.retrievers.append(vectorstore.as_retriever())
                            st.session_state.processed_files[file] = True
                        except Exception:
                            st.warning(f"Failed to load knowledge base: {file}. Skipping it.")

            # Step 5: File Upload (Admin only)
            if st.session_state.is_admin:
                uploaded_files = st.sidebar.file_uploader(
                    "Upload Files", type=['xlsx', 'xls', 'pdf', 'docx'], accept_multiple_files=True
                )

                if uploaded_files:
                    if not selected_category:
                        st.error("Please select a category before uploading files.")
                    else:
                        model_dir = os.path.join(DATA_DIR, selected_category, st.session_state.selected_model)
                        os.makedirs(model_dir, exist_ok=True)

                        existing_file_names = set(os.listdir(model_dir))
                        successful_uploads, failed_uploads = [], []

                        for uploaded_file in uploaded_files:
                            if uploaded_file.name in existing_file_names:
                                failed_uploads.append(uploaded_file.name)
                                continue

                            file_path = os.path.join(model_dir, uploaded_file.name)
                            with open(file_path, "wb") as f:
                                f.write(uploaded_file.getbuffer())
                            successful_uploads.append(uploaded_file)  
                            st.session_state.uploaded_file_names.append(uploaded_file.name)

                        if successful_uploads:
                            try:
                                process_uploaded_files(
                                    successful_uploads,  
                                    selected_category,
                                    st.session_state.selected_model
                                )
                                st.success(f"Uploaded and processed: {', '.join([file.name for file in successful_uploads])}")
                            except Exception as e:
                                st.error(f"Error processing files: {e}")

                        if failed_uploads:
                            st.warning(f"These files already exist and were skipped: {', '.join(failed_uploads)}")
                            
            # Step 6: Initialize QA Chain
            if st.session_state.retrievers:
                if len(st.session_state.retrievers) > 1:
                    ensemble_retriever = EnsembleRetriever(retrievers=st.session_state.retrievers)
                else:
                    ensemble_retriever = st.session_state.retrievers[0]

                try:
                    st.session_state.qa_chain = RetrievalQA.from_chain_type(
                        llm=st.session_state.llm,
                        chain_type='stuff',
                        retriever=ensemble_retriever,
                        verbose=True,
                    )
                except Exception:
                    st.error("Failed to initialize the QA chain. Please check the configuration.")
            # else:
                # st.error("No retrievers available. Please upload files or select a valid knowledge base.")

        except Exception as e:
            st.error(f"An unexpected error occurred: {str(e)}")

        # Chat container
        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []

        def render_chat_history():
            chat_history_html = "<div class='chat-history' style='height: 450px; overflow-y: auto; padding-right: 10px; padding-left: 10px; padding-top: 10px; padding-bottom: 10px;'>"
            for message in st.session_state.chat_history:
                if message['role'] == 'user':
                    chat_history_html += f"<div style='text-align: right;'><p style='color: #ffffff; background-color: #4caf50; padding: 10px; border-radius: 12px; max-width: 60%; display: inline-block;'>{message['message']}</p></div>"
                else:
                    chat_history_html += f"<div style='text-align: left;'><p style='color: #ffffff; background-color: #1e88e5; padding: 10px; border-radius: 12px; max-width: 60%; display: inline-block;'>{message['message']}</p></div>"
            chat_history_html += "</div>"
            return chat_history_html

        # Chat container
        chat_container = st.empty()

        # Input container
        input_container = st.container()
        with input_container:
            st.markdown("""
                <style>
                .stChatInput {
                    position: fixed;
                    bottom: 60px;
                    left: 50%;
                    transform: translateX(-50%);
                    width: 90%;
                    max-width: 600px;
                    z-index: 1000;
                }
                .stChatInput input {
                    width: 100%;
                    padding: 12px;
                    border: 2px solid #1e88e5;
                    border-radius: 12px;
                }
                </style>
            """, unsafe_allow_html=True)

            # User input field
            user_input = st.chat_input("Ask anything...")

            # Process user input
            if user_input:
                if user_input.strip():
                    # Append user's message to chat history and render immediately
                    user_message = {"role": "user", "message": user_input}
                    st.session_state.chat_history.append(user_message)
                    chat_container.markdown(render_chat_history(), unsafe_allow_html=True)

                    # Generate assistant's response
                    try:
                        with st.spinner("Thinking..."):
                            if st.session_state.qa_chain is None:
                                raise ValueError("QA chain is not initialized or is invalid.")

                            # Get the response from the QA chain
                            result = st.session_state.qa_chain({"query": user_input})
                            assistant_message = {"role": "assistant", "message": result['result']}
                            st.session_state.chat_history.append(assistant_message)

                            # Save to MongoDB
                            save_to_mongo(user_input, result['result'])

                    except ValueError as ve:
                        st.error(str(ve))
                    except Exception as e:
                        st.error(f"An unexpected error occurred: {str(e)}")

                    # Render updated chat history
                    chat_container.markdown(render_chat_history(), unsafe_allow_html=True)
   
                           
                    # if st.session_state.chat_history:
                    #     if st.button("Clear Chat"):
                    #         st.session_state.chat_history.clear()
                    #         st.session_state.clear_button_clicked = True
                    # else:
                    #     st.session_state.clear_button_clicked = False
                        
                    # if st.session_state.clear_button_clicked:
                    #     st.rerun()
                    
                
            #  User wise logged in data retrieval
            if 'show_date_fields' not in st.session_state:
                st.session_state.show_date_fields = False
            if st.sidebar.button("View Chat History"):
                st.session_state.show_date_fields = True
            if st.session_state.show_date_fields:
                start_date = st.sidebar.date_input("Choose Start Date", datetime.now())
                end_date = st.sidebar.date_input("Choose End Date", datetime.now())
                # Button to fetch chat history
                if st.sidebar.button("OK"):
                    # Validate the date range
                    if end_date < start_date:
                        st.sidebar.error("End date must be after the start date.")
                    else: 
                        chat_history = get_user_chat_history(st.session_state.username, start_date, end_date)

                        if chat_history:
                            st.sidebar.write("### Your Chat History:")
                            for entry in chat_history:
                                st.sidebar.write(f"**You**: {entry['user_query']}")
                                st.sidebar.write(f"**Bot**: {entry['assistant_response']}")
                                st.sidebar.write(f"_Timestamp: {entry['timestamp']}_")        
                        else:
                            st.sidebar.write("No chat history found.")
                        
    with st.sidebar:
        selected_option = option_menu(
            menu_title="",  
            options=["Home", "Chat","Models" ,"Filter Chat", "Logout"],  
            icons=["house", "chat-left-text", "box" , "chat-dots", "box-arrow-right"],  
            menu_icon="cast",  
            default_index=0,  
            orientation="vertical", 
            styles={
                "container": {
                    "padding": "0!important",
                    "background-color": "#ffffff", 
                    # "margin-top": "40px",
                    "position" : "relative",
                    "top":"20px",
                },
                "icon": {
                    "color": "#1a73e8",  
                    "font-size": "16px",  
                    "margin-right": "10px",  
                },
                "nav-link": {
                    "font-size": "16px",  
                      
                    "color": "#000000",  
                },
                "nav-link-selected": {
                    "background-color": "#f0f0f0",  
                    "color": "#1a73e8",  
                },
            },
        )
        
    # Logout 
    def logout():
        st.session_state.logged_in = False
        st.session_state.username = ""
        st.session_state.is_admin = False
        st.session_state.show_signup = False
        st.success("You have been logged out.")
        st.rerun()
             
    # Display 
    if selected_option == "Home":
        st.markdown(
        """
        <style>
        /* Full height layout with horizontal alignment */
        body, html {
            overflow: hidden;  /* Disable scrolling */
            margin: 0;  /* Remove default margin */
            padding: 0;  /* Remove default padding */
            font-family: Arial, sans-serif;  /* Smooth font */
            background-color: #ffffff;  /* White background for the body */
            display: flex
        }

        /* Sidebar styles */
        .sidebar {
            width: 300px;  /* Fixed width for sidebar */
            height: 100vh;  /* Full height */
            background-color: #f8f9fa;  /* Light gray background for sidebar */
            position: fixed;  /* Fixed position */
            left: 0;  /* Align to left */
            top: 0;  /* Align to top */
            padding: 20px;  /* Padding for content */
            box-shadow: 2px 0px 5px rgba(0, 0, 0, 0.1);  /* Shadow for depth */
        }

        /* Main content styles */
        .main-content {
            width:100%;
            padding: 20px;  /* Padding for main content */
            display: flex;  /* Flexbox for horizontal alignment */
            align-items: center;  /* Center vertically */
            justify-content: center;  /* Center horizontally */
        }

        /* Welcome box with smooth design */
        .welcome-box {
            background-color: white;  /* White box */
            border-radius: 15px;
            padding: 40px;
            max-width: 800px;
            width: 100%;
            text-align: center;
            transition: transform 0.3s ease, box-shadow 0.3s ease;
        }
        .welcome-box:hover {
            transform: translateY(-5px);  /* Soft hover effect */
            box-shadow: 0px 12px 25px rgba(0, 0, 0, 0.2);  /* Enhance shadow on hover */
        }

        /* Welcome header */
        .welcome-header {
            font-size: 2.5em;
            font-weight: bold;
            margin-bottom: 20px;
            color: #1a73e8;  /* Match navbar icon color */
        }

        /* Welcome message */
        .welcome-message {
            font-size: 1.2em;
            line-height: 1.6;
            color: #555;  /* Slightly lighter gray for readability */
        }

        /* Highlight key text */
        .highlight {
            color: #1a73e8;  /* Bright color for highlights, matching navbar */
            font-weight: bold;
        }

        /* Responsive design */
        @media (max-width: 768px) {
            .welcome-header {
                font-size: 2em;
            }
            .welcome-message {
                font-size: 1em;
            }
            .sidebar {
                width: 200px;  /* Adjust sidebar width for smaller screens */
            }
            .main-content {
                margin-left: 210px;  /* Adjust margin for smaller screens */
            }
        }
        </style>
        """,
        unsafe_allow_html=True
    )

        # Main content area for welcome message
        st.markdown(
    """
    <head>
        <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0-beta3/css/all.min.css">
    </head>
    <div class='main-content'>
        <div class='welcome-box'>
            <div class='welcome-header'>
                Welcome, {username}!
            </div>
            <div class='welcome-message'>
                <i class="fas fa-hand-spock"></i> We're excited to have you here! Ready to dive into some AI-powered conversations?<br><br>
                Simply <span class='highlight'>click on Chat </span> in the menu to start interacting with our intelligent models.<br><br>
                <i class="fas fa-search"></i> Want to explore specific topics? Select a <span class='highlight'>Domain</span> and <span class='highlight'>upload your file</span> and watch the AI bring your data to life!<br><br>
                Let the conversation begin! <i class="fas fa-rocket"></i>
            </div>
        </div>
    </div>
    """.format(username=st.session_state.username.capitalize()),
    unsafe_allow_html=True
)
            
    elif selected_option =="Filter Chat":
        load_chat_retrieval()
        
    elif selected_option == "Chat":
        initiate_chat()
        
    elif selected_option == "Models":
        st.markdown(
    """
    <style>
    .subheader {
        font-size: 24px;
        font-weight: bold;
        color: #007bff; /* Primary blue color */
        margin-bottom: 10px;
        text-align: center;
    }

    /* Styling for the descriptive text */
    .description {
        font-size: 18px;
        color: #333333; /* Dark grey for readability */
        background-color: #f9f9f9; /* Light grey background */
        border-left: 4px solid #66bb6a; /* Green border to highlight */
        padding: 15px;
        border-radius: 6px;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1); /* Light shadow for depth */
        margin-bottom: 20px;
        line-height: 1.5;
    }
    /* General styling for the DataFrame table */
    .model-table {
        width: 100%;
        border-collapse: collapse;
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 4px 10px rgba(0, 0, 0, 0.1);
    }

    .model-table th, .model-table td {
        padding: 12px;
        border: 1px solid #e0e0e0;
        text-align: left;
        font-size: 16px;
    }

    .model-table thead {
        background-color: #007bff;
        color: white;
    }

    .model-table tbody tr:nth-child(even) {
        background-color: #f9f9f9;
    }

    .model-table tbody tr:nth-child(odd) {
        background-color: #ffffff;
    }

    .model-table tbody tr:hover {
        background-color: #e8f5e9;
    }
    
    </style>
    """,
    unsafe_allow_html=True
)   
        available_models = list_models(MODEL_DIR)
        if available_models:
            st.markdown('<div class="subheader">Unlock the Power of Advanced LLM Models for Rich Conversations:</div>', unsafe_allow_html=True)

            st.markdown("""
    <div class="description">
        Discover the next generation of intelligent conversation with our advanced Large Language Models (LLMs). These innovative models are designed to enhance your dialogue experience, providing profound insights, nuanced replies, and fluid interactions that closely mimic human conversation. Browse the models below to embark on engaging and personalized discussions tailored just for you.
    </div>
""", unsafe_allow_html=True)
            
            # Create a DataFrame to hold the model names with serial numbers
            model_data = pd.DataFrame({
                "S.No.": range(1, len(available_models) + 1),
                "Model Name": available_models
            })
            
            # Render the table using styled HTML
            st.markdown(model_data.to_html(index=False, classes="model-table"), unsafe_allow_html=True)
        else:
            st.warning("⚠️ No models available")
            
    elif selected_option == "Logout":
        logout()
